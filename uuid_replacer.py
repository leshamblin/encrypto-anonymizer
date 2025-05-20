#!/usr/bin/env python3
"""
UUID Replacer for DOCX files
This script reads UUIDs from anonymized.csv and replaces them with original names
from the mapping.json file in DOCX documents.
"""

import json
import csv
import argparse
from pathlib import Path
from docx import Document
from cryptography.fernet import Fernet


def load_key(key_path='key.key'):
    """Load the encryption key from file"""
    with open(key_path, 'rb') as key_file:
        return key_file.read()


def decrypt_value(encrypted_value, fernet):
    """Decrypt an encrypted value using Fernet"""
    if not encrypted_value:
        return encrypted_value
    try:
        return fernet.decrypt(encrypted_value.encode()).decode()
    except Exception as e:
        print(f"Error decrypting value: {e}")
        return encrypted_value


def load_mapping(mapping_path='mapping.json', key_path='key.key'):
    """Load and decrypt the mapping from JSON file"""
    key = load_key(key_path)
    fernet = Fernet(key)
    
    with open(mapping_path, 'r') as file:
        encrypted_mapping = json.load(file)
    
    # Check if encrypted_mapping is a dictionary
    if not isinstance(encrypted_mapping, dict):
        raise TypeError(f"Expected dictionary from {mapping_path}, got {type(encrypted_mapping).__name__}")
        
    # Decrypt the mapping
    mapping = {}
    for uuid, encrypted_data in encrypted_mapping.items():
        # Check if encrypted_data is a dictionary or string
        if isinstance(encrypted_data, dict):
            decrypted_data = {}
            for field, encrypted_value in encrypted_data.items():
                decrypted_data[field] = decrypt_value(encrypted_value, fernet)
            mapping[uuid] = decrypted_data
        elif isinstance(encrypted_data, str):
            # Handle case where encrypted_data is a string
            mapping[uuid] = decrypt_value(encrypted_data, fernet)
        else:
            print(f"Warning: Unexpected data type for UUID {uuid}: {type(encrypted_data).__name__}")
    
    return mapping


def get_uuids_from_csv(csv_path='anonymized.csv'):
    """Extract unique UUIDs and column fields from the anonymized CSV file"""
    uuids = set()
    columns = []
    uuid_columns = {}  # Maps UUIDs to the column they were found in
    
    with open(csv_path, 'r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        columns = reader.fieldnames or []
        
        for row in reader:
            # Check all columns for UUID-like values
            for column, value in row.items():
                if value and '-' in value and len(value) == 36:
                    # Basic UUID format check (8-4-4-4-12 pattern)
                    try:
                        parts = value.split('-')
                        if len(parts) == 5 and len(parts[0]) == 8 and len(parts[1]) == 4 and \
                           len(parts[2]) == 4 and len(parts[3]) == 4 and len(parts[4]) == 12:
                            uuids.add(value)
                            # Track which column this UUID was found in
                            if value not in uuid_columns:
                                uuid_columns[value] = []
                            if column not in uuid_columns[value]:
                                uuid_columns[value].append(column)
                    except:
                        pass
    
    return uuids, columns, uuid_columns


def replace_uuids_in_paragraph(paragraph, uuid_to_name):
    """Replace UUIDs in a paragraph with original names"""
    text = paragraph.text
    
    for uuid, name in uuid_to_name.items():
        if uuid in text:
            text = text.replace(uuid, name)
    
    if text != paragraph.text:
        # Clear the paragraph and add the new text
        paragraph.clear()
        paragraph.add_run(text)


def replace_uuids_in_table(table, uuid_to_name):
    """Replace UUIDs in table cells with original names"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_uuids_in_paragraph(paragraph, uuid_to_name)


def process_docx(docx_path, uuid_to_name, output_path=None):
    """Process a DOCX file and replace UUIDs with names"""
    doc = Document(docx_path)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_uuids_in_paragraph(paragraph, uuid_to_name)
    
    # Replace in tables
    for table in doc.tables:
        replace_uuids_in_table(table, uuid_to_name)
    
    # Save the document
    if output_path is None:
        output_path = docx_path.stem + '_deanonymized.docx'
    
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Replace UUIDs with original names in DOCX files')
    parser.add_argument('docx_file', help='Path to the DOCX file to process')
    parser.add_argument('-o', '--output', help='Output file path (default: input_deanonymized.docx)')
    parser.add_argument('-m', '--mapping', default='mapping.json', help='Path to mapping.json')
    parser.add_argument('-c', '--csv', default='anonymized.csv', help='Path to anonymized.csv')
    parser.add_argument('-k', '--key', default='key.key', help='Path to key.key')
    
    args = parser.parse_args()
    
    # Load the mapping
    print("Loading mapping...")
    mapping = load_mapping(args.mapping, args.key)
    
    # Get UUIDs from CSV
    print("Getting UUIDs from CSV...")
    uuids, csv_columns, uuid_columns = get_uuids_from_csv(args.csv)
    
    print(f"Found {len(csv_columns)} columns in CSV: {', '.join(csv_columns)}")
    
    # Create UUID to name mapping
    uuid_to_name = {}
    field_mapping = {}  # Keep track of which field each UUID came from
    
    # First, identify which field each UUID belongs to in the mapping
    for uuid, data in mapping.items():
        # Check if data is a dictionary before trying to iterate through its items
        if isinstance(data, dict):
            for field, value in data.items():
                if value == uuid:
                    field_mapping[uuid] = field
    
    # Now create the UUID to name mapping
    for uuid in uuids:
        if uuid in mapping:
            # Get the decrypted original value for each field
            data = mapping[uuid]
            
            # Check if data is a string 
            if isinstance(data, str):
                uuid_to_name[uuid] = data
            # Otherwise treat as dictionary
            elif isinstance(data, dict):
                # Get columns where this UUID was found
                columns_for_uuid = uuid_columns.get(uuid, [])
                
                # First, try to use a field that matches one of the columns in the CSV
                found_match = False
                for column in columns_for_uuid:
                    # Try to find this column name in the data
                    if column in data:
                        uuid_to_name[uuid] = data[column]
                        found_match = True
                        break
                
                # If no direct column match, use field_mapping if available
                if not found_match and uuid in field_mapping:
                    field = field_mapping[uuid]
                    if field in data:
                        uuid_to_name[uuid] = data[field]
                        found_match = True
                
                # If still no match, use any available field (try all CSV columns first)
                if not found_match:
                    for column in csv_columns:
                        if column in data and data[column]:
                            uuid_to_name[uuid] = data[column]
                            found_match = True
                            break
                    
                    # If still no match, use any field with a value
                    if not found_match:
                        for field, value in data.items():
                            if value:
                                uuid_to_name[uuid] = value
                                found_match = True
                                break
    
    print(f"Found {len(uuid_to_name)} UUIDs to replace")
    
    # Process the DOCX file
    print(f"Processing {args.docx_file}...")
    output_path = process_docx(args.docx_file, uuid_to_name, args.output)
    
    print(f"Done! Output saved to: {output_path}")


if __name__ == "__main__":
    main()